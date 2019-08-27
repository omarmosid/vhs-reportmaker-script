import xlrd
file_location = "F:/Learning/Python/reportmakerapp/Excel/Class X.xlsx"
workbook = xlrd.open_workbook(file_location)
sheet = workbook.sheet_by_index(0)

from docx import Document
document = Document('template.docx')

for row in range(3, sheet.nrows):
    for col in range(sheet.ncols):
        #Student Name
        if(col == 1):
            print("Student Name " + sheet.cell_value(row, col))
            document.paragraphs[0].runs[2].text = sheet.cell_value(row, col)
        #Conceptual Understanding
        if(col == 3):
            print("Conceptual Understanding " + sheet.cell_value(row, col))
            document.paragraphs[5].runs[1].text = sheet.cell_value(row, col)
        
        #Listening Skills
        if(col == 4):
            print("Listening Skills " + sheet.cell_value(row, col))
            document.paragraphs[11].runs[2].text = sheet.cell_value(row, col)
        
        #Speaking Skills
        if(col == 5):
            print("Speaking Skills " + sheet.cell_value(row, col))
            document.paragraphs[13].runs[2].text = sheet.cell_value(row, col)
        
        #Reading Skills
        if(col == 6):
            print("Reading Skills " + sheet.cell_value(row, col))
            document.paragraphs[16].runs[2].text = sheet.cell_value(row, col)
        
        #Writing Skills
        #Organisation and content
        if(col == 7):
            print("Organisation and content " + sheet.cell_value(row, col))
            document.paragraphs[20].runs[1].text = sheet.cell_value(row, col)
        
        #Presentation and submission
        if(col == 8):
            print("Presentation and submission " + sheet.cell_value(row, col))
            document.paragraphs[23].runs[2].text = sheet.cell_value(row, col)
        
        #Vocabulary and grammar
        if(col == 9):
            print("Vocabulary and grammar " + sheet.cell_value(row, col))
            document.paragraphs[26].runs[2].text = sheet.cell_value(row, col)
        
        #Classroom involvement
        if(col == 10):
            print("Classroom involvement " + sheet.cell_value(row, col))
            document.paragraphs[29].runs[1].text = sheet.cell_value(row, col)
        
        #Test performance
        if(col == 11):
            print("Test performance " + sheet.cell_value(row, col))
            document.paragraphs[32].runs[1].text = sheet.cell_value(row, col)
        
        #Suggestions regarding areas/skills to be improved
        if(col == 12):
            print("Suggestions regarding areas/skills to be improved " + sheet.cell_value(row, col))
            document.paragraphs[35].runs[2].text = sheet.cell_value(row, col)
        
        #Comments
        if(col == 13):
            print("Comments " + sheet.cell_value(row, col))
            document.paragraphs[38].runs[1].text = sheet.cell_value(row, col)
    
    #save file using student name with docx extension
    document.save('F:/Learning/Python/reportmakerapp/exported/' + sheet.cell_value(row, 1) + '.docx')

print("")
print("Done!!!!!!!!")

# document.save('test.docx')